import copy
import os
import re
from datetime import datetime
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
from xml.sax.saxutils import escape, unescape
import xml.etree.ElementTree as ET
from decimal import Decimal

from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect
from django.template.loader import get_template
from django.shortcuts import render
from django.urls import reverse_lazy
from django.views.generic.base import View
from weasyprint import HTML, CSS

from config import settings
from core.pos.brand_assets import comprobante_logo_url, comprobante_print_context
from core.pos.client_properties import lock_client_properties_for_sale_contract
from core.pos.models import Sale, Company, CtasCollect
from core.pos.views.crm.sale.print.contract_docx_decorate import decorate_contract_celia_docx
from core.pos.views.crm.sale.print.docx_codes import decorate_contract_docx


def _contract_code_for_documents(sale):
    return (sale.contract_code or f'CT-{sale.pk:06d}').strip()


def _company_website_text():
    c = Company.objects.first()
    return (c.website or '').strip() if c else ''


def _company_description_text():
    """Descripción corta para el texto bajo el QR del contrato (campo desc en Empresa)."""
    c = Company.objects.first()
    return (c.desc or '').strip() if c else ''


def _client_code_for_sale(sale):
    client = sale.client
    if not client:
        return ''
    return (client.client_code or '').strip()


def _sale_code_for_sale(sale):
    return (sale.sale_code or f'VT-{sale.pk:06d}').strip()


def _is_celia_contract_template(template_path):
    if not template_path:
        return False
    low = os.path.basename(template_path).lower()
    return 'celia' in low or 'notarial' in low


def _build_decorated_contract_docx(template_bytes, sale, mapping, signature_date_plain, template_path=None):
    docx_out = _apply_mapping_to_docx(
        template_bytes,
        mapping,
        sale=sale,
        signature_date_plain=signature_date_plain,
        location_ctx=_contract_location_ctx(sale) if _is_celia_contract_template(template_path) else None,
        community=_contract_linked_property_community(sale),
    )
    if _is_celia_contract_template(template_path):
        return decorate_contract_celia_docx(
            docx_out,
            _client_code_for_sale(sale),
            _sale_code_for_sale(sale),
            _contract_code_for_documents(sale),
        )
    return decorate_contract_docx(
        docx_out,
        _contract_code_for_documents(sale),
        _company_website_text(),
        _company_description_text(),
    )


MONTHS_ES_UPPER = {
    1: 'ENERO',
    2: 'FEBRERO',
    3: 'MARZO',
    4: 'ABRIL',
    5: 'MAYO',
    6: 'JUNIO',
    7: 'JULIO',
    8: 'AGOSTO',
    9: 'SEPTIEMBRE',
    10: 'OCTUBRE',
    11: 'NOVIEMBRE',
    12: 'DICIEMBRE',
}

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W_P = '{%s}p' % W_NS
W_R = '{%s}r' % W_NS
W_T = '{%s}t' % W_NS
W_PPR = '{%s}pPr' % W_NS
W_RPR = '{%s}rPr' % W_NS
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

_MONTH_NAMES_RE = (
    'enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre|'
    'ENERO|FEBRERO|MARZO|ABRIL|MAYO|JUNIO|JULIO|AGOSTO|SEPTIEMBRE|OCTUBRE|NOVIEMBRE|DICIEMBRE'
)
_SIGNATURE_FIRMA_DATE_RE = re.compile(
    rf'(,\s*)?(\d{{1,2}}\s+de\s+(?:{_MONTH_NAMES_RE})\s+del\s+\d{{4}}\.?)',
    re.IGNORECASE,
)


def _contract_template_path():
    folder = os.path.join(
        settings.BASE_DIR,
        'core', 'pos', 'templates', 'contracts',
    )
    for name in (
        'CONTRATO NOTARIAL - FINAL.docx',
        'contrato_Celia.docx',
        'CONTRATO SRA. CELIA.docx',
        'contrato_base.docx',
    ):
        path = os.path.join(folder, name)
        if os.path.isfile(path):
            return path
    return os.path.join(folder, 'contrato_Celia.docx')


def _schedule_template_path():
    return os.path.join(
        settings.BASE_DIR,
        'core', 'pos', 'templates', 'contracts', 'Cronograma.docx',
    )


def _ajax_request(request):
    return request.headers.get('X-Requested-With') == 'XMLHttpRequest'


def _contract_download_error_response(request, message, status=500):
    if _ajax_request(request):
        return HttpResponse(message, status=status, content_type='text/plain; charset=utf-8')
    try:
        if (
            getattr(request, 'user', None)
            and request.user.is_authenticated
            and request.user.is_client()
        ):
            return HttpResponseRedirect(reverse_lazy('sale_client_list'))
    except Exception:
        pass
    return HttpResponseRedirect(reverse_lazy('sale_admin_list'))


def _contract_signature_date_plain(sale):
    sale_day = sale.date_joined
    contract_date = sale_day.date() if hasattr(sale_day, 'date') else sale_day
    return (
        f'{contract_date.day} de {MONTHS_ES_UPPER.get(contract_date.month, "")} '
        f'del {contract_date.year}.'
    )


def _contract_linked_property(sale):
    detail = (
        sale.saledetail_set
        .select_related('client_property')
        .filter(client_property__isnull=False)
        .order_by('id')
        .first()
    )
    return detail.client_property if detail else None


def _contract_linked_property_community(sale):
    prop = _contract_linked_property(sale)
    if prop and (prop.community or '').strip():
        return prop.community.strip().upper()
    return ''


def _text(value, fallback=''):
    value = '' if value is None else str(value).strip()
    return value or fallback


def _upper(value, fallback=''):
    return _text(value, fallback).upper()


def _date_iso(value):
    if not value:
        return ''
    value = value.date() if hasattr(value, 'date') else value
    return value.strftime('%Y-%m-%d')


def _property_predio_type(prop):
    if not prop or not prop.predio_type:
        return ''
    try:
        return prop.get_predio_type_display()
    except Exception:
        return prop.predio_type


def _client_predio_type(client):
    if not client or not client.predio_type:
        return ''
    try:
        return client.get_predio_type_display()
    except Exception:
        return client.predio_type


def _predio_value(prop, client, prop_attr, client_attr='', fallback=''):
    if prop and _text(getattr(prop, prop_attr, '')):
        return _text(getattr(prop, prop_attr, ''))
    if client_attr and client and _text(getattr(client, client_attr, '')):
        return _text(getattr(client, client_attr, ''))
    return fallback


def _sale_products_text(sale):
    names = []
    for detail in sale.saledetail_set.select_related('product').order_by('id'):
        if detail.product and detail.product.name:
            names.append(detail.product.name.strip())
    return ', '.join(names)


def _money_amounts_for_sale(sale):
    total_amount = Decimal(str(sale.total or 0)).quantize(Decimal('0.01'))
    down_payment = Decimal('0.00')
    if sale.payment_condition == 'credito':
        down_payment = Decimal(str(sale.credit_down_payment or 0)).quantize(Decimal('0.01'))
    remaining_amount = (total_amount - down_payment).quantize(Decimal('0.01'))
    if remaining_amount < 0:
        remaining_amount = Decimal('0.00')
    return total_amount, down_payment, remaining_amount


def _format_money(amount):
    return 'S/ {:,.2f}'.format(Decimal(str(amount)).quantize(Decimal('0.01')))


def _number_to_words_es(num):
    num = int(num)
    units = (
        'CERO', 'UNO', 'DOS', 'TRES', 'CUATRO', 'CINCO', 'SEIS', 'SIETE', 'OCHO', 'NUEVE',
        'DIEZ', 'ONCE', 'DOCE', 'TRECE', 'CATORCE', 'QUINCE', 'DIECISEIS', 'DIECISIETE',
        'DIECIOCHO', 'DIECINUEVE', 'VEINTE', 'VEINTIUNO', 'VEINTIDOS', 'VEINTITRES',
        'VEINTICUATRO', 'VEINTICINCO', 'VEINTISEIS', 'VEINTISIETE', 'VEINTIOCHO',
        'VEINTINUEVE',
    )
    tens = {
        30: 'TREINTA',
        40: 'CUARENTA',
        50: 'CINCUENTA',
        60: 'SESENTA',
        70: 'SETENTA',
        80: 'OCHENTA',
        90: 'NOVENTA',
    }
    hundreds = {
        100: 'CIEN',
        200: 'DOSCIENTOS',
        300: 'TRESCIENTOS',
        400: 'CUATROCIENTOS',
        500: 'QUINIENTOS',
        600: 'SEISCIENTOS',
        700: 'SETECIENTOS',
        800: 'OCHOCIENTOS',
        900: 'NOVECIENTOS',
    }
    if num < 30:
        return units[num]
    if num < 100:
        base = (num // 10) * 10
        rest = num % 10
        return tens[base] if rest == 0 else '{} Y {}'.format(tens[base], units[rest])
    if num < 1000:
        if num in hundreds:
            return hundreds[num]
        base = (num // 100) * 100
        rest = num % 100
        prefix = 'CIENTO' if base == 100 else hundreds[base]
        return '{} {}'.format(prefix, _number_to_words_es(rest))
    if num < 1000000:
        thousands = num // 1000
        rest = num % 1000
        prefix = 'MIL' if thousands == 1 else '{} MIL'.format(_number_to_words_es(thousands))
        return prefix if rest == 0 else '{} {}'.format(prefix, _number_to_words_es(rest))
    return str(num)


def _amount_to_contract_words(amount):
    amount = Decimal(str(amount)).quantize(Decimal('0.01'))
    whole = int(amount)
    cents = int((amount - Decimal(whole)) * 100)
    return '{} Y {:02d}/100 SOLES'.format(_number_to_words_es(whole), cents)


def _contract_placeholder_mapping(sale):
    client = sale.client
    prop = _contract_linked_property(sale)
    user = client.user if client else None
    company = Company.objects.first()

    total_amount, down_payment, remaining_amount = _money_amounts_for_sale(sale)

    client_department = _upper(client.department if client else '')
    client_province = _upper(client.province if client else '')
    client_district = _upper(client.district if client else '')
    predio_department = _upper(_predio_value(prop, client, 'department', 'predio_department'))
    predio_province = _upper(_predio_value(prop, client, 'province', 'predio_province'))
    predio_district = _upper(_predio_value(prop, client, 'district', 'predio_district'))
    predio_address = _upper(_predio_value(prop, client, 'address', 'predio_address'))
    predio_lot = _upper(_predio_value(prop, client, 'lot_number', 'predio_lot_number'))
    predio_block = _upper(_predio_value(prop, client, 'block', 'predio_block'))
    predio_registry = _upper(_predio_value(prop, client, 'registry_number', 'predio_registry_number'))
    predio_area = (
        format(prop.area, '.2f') if prop and prop.area is not None
        else format(client.predio_area, '.2f') if client and client.predio_area is not None
        else ''
    )
    predio_perimeter = (
        format(prop.perimeter, '.2f') if prop and prop.perimeter is not None
        else format(client.predio_perimeter, '.2f') if client and client.predio_perimeter is not None
        else ''
    )
    predio_type = _upper(_property_predio_type(prop) or _client_predio_type(client))
    predio_community = _upper(prop.community if prop else '')
    predio_population_center = _upper(prop.population_center if prop else '')
    sale_products = _upper(_sale_products_text(sale))
    employee_name = _upper(sale.employee.get_full_name() if sale.employee else '')

    try:
        payment_condition = sale.get_payment_condition_display()
    except Exception:
        payment_condition = sale.payment_condition or ''
    try:
        payment_method = sale.get_payment_method_display()
    except Exception:
        payment_method = sale.payment_method or ''

    values = {
        'CLIENTE_NOMBRE': _upper(user.get_full_name() if user else '', 'CLIENTE'),
        'NOMBRE_CLIENTE': _upper(user.get_full_name() if user else '', 'CLIENTE'),
        'CLIENTE_DNI': _text(user.dni if user else '', '00000000'),
        'DNI_CLIENTE': _text(user.dni if user else '', '00000000'),
        'CLIENTE_CELULAR': _text(client.mobile if client else '', '000000000'),
        'CELULAR_CLIENTE': _text(client.mobile if client else '', '000000000'),
        'CLIENTE_ESTADO_CIVIL': _upper(
            client.get_marital_status_display() if client and client.marital_status else '',
            '—',
        ),
        'ESTADO_CIVIL': _upper(
            client.get_marital_status_display() if client and client.marital_status else '',
            '—',
        ),
        'CLIENTE_DIRECCION': _upper(client.address if client else '', '—'),
        'DIRECCION_CLIENTE': _upper(client.address if client else '', '—'),
        'CLIENTE_DEPARTAMENTO': client_department,
        'DEPARTAMENTO_CLIENTE': client_department,
        'CLIENTE_PROVINCIA': client_province,
        'PROVINCIA_CLIENTE': client_province,
        'CLIENTE_DISTRITO': client_district,
        'DISTRITO_CLIENTE': client_district,
        'CODIGO_CLIENTE': _client_code_for_sale(sale),
        'CODIGO_VENTA': _sale_code_for_sale(sale),
        'CODIGO_CONTRATO': _contract_code_for_documents(sale),
        'PREDIO_COMUNIDAD': predio_community,
        'COMUNIDAD_PREDIO': predio_community,
        'COMUNIDAD': predio_community,
        'PREDIO_CENTRO_POBLADO': predio_population_center,
        'CENTRO_POBLADO': predio_population_center,
        'PREDIO_DEPARTAMENTO': predio_department,
        'DEPARTAMENTO_PREDIO': predio_department,
        'PREDIO_PROVINCIA': predio_province,
        'PROVINCIA_PREDIO': predio_province,
        'PREDIO_DISTRITO': predio_district,
        'DISTRITO_PREDIO': predio_district,
        'PREDIO_DIRECCION': predio_address,
        'DIRECCION_PREDIO': predio_address,
        'PREDIO_LOTE': predio_lot,
        'LOTE_PREDIO': predio_lot,
        'PREDIO_MANZANA': predio_block,
        'MANZANA_PREDIO': predio_block,
        'PREDIO_PARTIDA': predio_registry,
        'PARTIDA_PREDIO': predio_registry,
        'PREDIO_TIPO': predio_type,
        'TIPO_PREDIO': predio_type,
        'PREDIO_AREA': predio_area,
        'AREA_PREDIO': predio_area,
        'PREDIO_PERIMETRO': predio_perimeter,
        'PERIMETRO_PREDIO': predio_perimeter,
        'SERVICIO': sale_products,
        'PRODUCTO': sale_products,
        'PRODUCTOS': sale_products,
        'ASESOR': employee_name,
        'VENDEDOR': employee_name,
        'FORMA_PAGO': _upper(payment_condition),
        'CONDICION_PAGO': _upper(payment_condition),
        'METODO_PAGO': _upper(payment_method),
        'EMPRESA_NOMBRE': _upper(company.name if company else ''),
        'EMPRESA_RUC': _text(company.ruc if company else ''),
        'EMPRESA_DIRECCION': _upper(company.address if company else ''),
        'EMPRESA_CELULAR': _text(company.mobile if company else ''),
        'EMPRESA_TELEFONO': _text(company.phone if company else ''),
        'EMPRESA_EMAIL': _text(company.email if company else ''),
        'EMPRESA_WEB': _text(company.website if company else ''),
        'EMPRESA_DESCRIPCION': _text(company.desc if company else ''),
        'MONTO_TOTAL': _format_money(total_amount),
        'MONTO_TOTAL_LETRAS': _amount_to_contract_words(total_amount),
        'VENTA_TOTAL': _format_money(total_amount),
        'CREDITO_INICIAL': _format_money(down_payment),
        'INICIAL': _format_money(down_payment),
        'PRIMER_PAGO': _format_money(down_payment),
        'SALDO_CREDITO': _format_money(remaining_amount),
        'SALDO': _format_money(remaining_amount),
        'NUMERO_CUOTAS': str(sale.credit_quota_count or 1),
        'CUOTAS': str(sale.credit_quota_count or 1),
        'FECHA_VENTA': _date_iso(sale.date_joined),
        'FECHA_LIMITE_CREDITO': _date_iso(sale.end_credit),
        'FECHA_CONTRATO': _contract_signature_date_plain(sale).rstrip('.'),
        'FECHA_FIRMA': _contract_signature_date_plain(sale).rstrip('.'),
    }

    mapping = {}
    wrappers = (
        ('{{', '}}'),
        ('[[', ']]'),
        ('${', '}'),
        ('<<', '>>'),
        ('«', '»'),
    )
    for key, value in values.items():
        escaped = escape(value)
        for left, right in wrappers:
            mapping[f'{left}{key}{right}'] = escaped
    return mapping


def _paragraph_join_text(p_elem):
    parts = []
    for t_elem in p_elem.findall('.//{%s}t' % W_NS):
        if t_elem.text:
            parts.append(t_elem.text)
    return ''.join(parts)


def _clone_run_properties_from_paragraph(p_elem):
    """
    Copia w:rPr de un run del párrafo para no perder fuente (p. ej. Arial Narrow), tamaño ni negrita.
    Prioriza el run con el fragmento de texto más largo (suele ser el cuerpo de la frase).
    """
    runs = p_elem.findall('.//{%s}r' % W_NS)
    best_rpr = None
    best_len = -1
    for r in runs:
        t = r.find(W_T)
        ln = len(t.text) if t is not None and t.text else 0
        rpr = r.find(W_RPR)
        if rpr is not None and ln >= best_len:
            best_len = ln
            best_rpr = rpr
    if best_rpr is not None:
        return copy.deepcopy(best_rpr)
    for r in reversed(runs):
        rpr = r.find(W_RPR)
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _paragraph_set_plain_text(p_elem, new_text):
    r_pr_clone = _clone_run_properties_from_paragraph(p_elem)
    p_pr = None
    for child in list(p_elem):
        if child.tag == W_PPR:
            p_pr = child
        p_elem.remove(child)
    if p_pr is not None:
        p_elem.append(p_pr)
    r = ET.Element(W_R)
    if r_pr_clone is not None:
        r.append(r_pr_clone)
    t = ET.Element(W_T)
    t.set(XML_SPACE, 'preserve')
    t.text = new_text
    r.append(t)
    p_elem.append(r)


def _replace_signature_date_in_document_xml(xml_str, new_date_plain):
    """
    Word suele partir '14 de ABRIL del 2026' en varios <w:t>, por eso .replace() no funciona.
    Se localiza el párrafo (p. ej. con 'Huancavelica') y se sustituye la fecha en el texto unido.
    """
    if not new_date_plain:
        return xml_str
    new_date_plain = new_date_plain.strip()

    def _repl(m):
        prefix = m.group(1) if m.group(1) else ', '
        return prefix + new_date_plain

    try:
        root = ET.fromstring(xml_str.encode('utf-8'))
    except ET.ParseError:
        return xml_str

    def _try_paragraphs(predicate):
        for p_elem in root.iter(W_P):
            full = _paragraph_join_text(p_elem)
            if not predicate(full):
                continue
            new_full, n = _SIGNATURE_FIRMA_DATE_RE.subn(_repl, full, count=1)
            if n:
                _paragraph_set_plain_text(p_elem, new_full)
                return True
        return False

    if not _try_paragraphs(lambda s: 'huancavelica' in s.lower()):
        _try_paragraphs(
            lambda s: 'conformidad' in s.lower()
            and ('firman' in s.lower() or 'firma' in s.lower())
        )

    ET.register_namespace('w', W_NS)
    return ET.tostring(root, encoding='unicode', method='xml')


def _replace_contract_community_in_document_xml(xml_str, community):
    if not community:
        return xml_str
    community = community.strip().upper()
    if not community:
        return xml_str
    try:
        root = ET.fromstring(xml_str.encode('utf-8'))
    except ET.ParseError:
        return xml_str

    patterns = (
        re.compile(r'((?:comunidad|comunidad\s+campesina)\s+de\s+)([^,.;]+)([,.;])', re.IGNORECASE),
        re.compile(r'(Centro\s+Poblado\s+de\s+)([^,.;]+)([,.;])', re.IGNORECASE),
    )
    changed = False
    for p_elem in root.iter(W_P):
        full = _paragraph_join_text(p_elem)
        if 'comunidad' not in full.lower() and 'centro poblado' not in full.lower():
            continue
        new_full = full
        for pattern in patterns:
            new_full = pattern.sub(lambda m: m.group(1) + community + m.group(3), new_full)
        if new_full != full:
            _paragraph_set_plain_text(p_elem, new_full)
            changed = True
    if not changed:
        return xml_str
    ET.register_namespace('w', W_NS)
    return ET.tostring(root, encoding='unicode', method='xml')


def _replace_placeholders_in_xml_paragraphs(xml_str, mapping):
    if not mapping:
        return xml_str
    try:
        root = ET.fromstring(xml_str.encode('utf-8'))
    except ET.ParseError:
        return xml_str

    changed = False
    for p_elem in root.iter(W_P):
        full = _paragraph_join_text(p_elem)
        if not full:
            continue
        new_full = full
        for old, new in mapping.items():
            if old in new_full:
                new_full = new_full.replace(old, unescape(new))
        if new_full != full:
            _paragraph_set_plain_text(p_elem, new_full)
            changed = True

    if not changed:
        return xml_str
    ET.register_namespace('w', W_NS)
    return ET.tostring(root, encoding='unicode', method='xml')


def _replace_contract_payment_amounts_in_document_xml(xml_str, sale):
    total_amount, down_payment, _remaining_amount = _money_amounts_for_sale(sale)
    total_text = _format_money(total_amount)
    total_words = _amount_to_contract_words(total_amount)
    down_payment_text = _format_money(down_payment)
    try:
        root = ET.fromstring(xml_str.encode('utf-8'))
    except ET.ParseError:
        return xml_str

    changed = False
    money_re = re.compile(r'S/\s*[\d,]+(?:\.\d{2})?')
    total_words_re = re.compile(r'\(([A-ZÁÉÍÓÚÑ\s]+?\d{2}/100\s+SOLES)\)', re.IGNORECASE)

    for p_elem in root.iter(W_P):
        full = _paragraph_join_text(p_elem)
        if not full:
            continue

        normalized = full.lower()
        new_full = full
        if 'monto total' in normalized and 'servicios' in normalized:
            new_full = money_re.sub(total_text, new_full, count=1)
            new_full = total_words_re.sub('({})'.format(total_words), new_full, count=1)
        elif '1er' in normalized and 'pago' in normalized:
            new_full = money_re.sub(down_payment_text, new_full, count=1)

        if new_full != full:
            _paragraph_set_plain_text(p_elem, new_full)
            changed = True

    if not changed:
        return xml_str
    ET.register_namespace('w', W_NS)
    return ET.tostring(root, encoding='unicode', method='xml')


def _contract_location_ctx(sale):
    client = sale.client
    prop = _contract_linked_property(sale)
    predio_dep = (
        (prop.department if prop and prop.department else '')
        or (client.predio_department if client and client.predio_department else '')
        or (client.department if client else '')
        or 'HUANCAVELICA'
    )
    predio_prov = (
        (prop.province if prop and prop.province else '')
        or (client.predio_province if client and client.predio_province else '')
        or (client.province if client else '')
        or 'HUANCAVELICA'
    )
    predio_dist = (
        (prop.district if prop and prop.district else '')
        or (client.predio_district if client and client.predio_district else '')
        or (client.district if client else '')
        or 'HUANCAVELICA'
    )
    return {
        'district': predio_dist.strip().upper(),
        'province': predio_prov.strip().upper(),
        'department': predio_dep.strip().upper(),
    }


def _contract_signature_date_celia(sale):
    sale_day = sale.date_joined
    contract_date = sale_day.date() if hasattr(sale_day, 'date') else sale_day
    return (
        f'{contract_date.day} de {MONTHS_ES_UPPER.get(contract_date.month, "")} '
        f'del {contract_date.year}'
    )


def _contract_mapping_celia(sale):
    client = sale.client
    prop = _contract_linked_property(sale)
    user = client.user if client else None
    full_name = (user.get_full_name() if user else '').strip().upper() or 'CLIENTE'
    dni = (user.dni if user else '').strip() or '00000000'
    mobile = (client.mobile if client and client.mobile else '').strip() or '000000000'
    address = (
        (prop.address if prop and prop.address else '')
        or (client.predio_address if client and client.predio_address else '')
        or (client.address if client else '')
        or '—'
    ).strip().upper()
    total_amount = Decimal(str(sale.total or 0)).quantize(Decimal('0.01'))
    contract_date_text = _contract_signature_date_celia(sale)

    mapping = {
        'CELIA JUSTINA BUSSO ÑAHUI': escape(full_name),
        '23260963': escape(dni),
        '971895287': escape(mobile),
        'Av. C. Manchego Muñoz N°1162 - BARRIO DE SANTA ANA': escape(address),
        'MONTO TOTAL S/3000.00': escape(f'MONTO TOTAL S/{total_amount:.2f}'),
        'S/3000.00': escape(f'S/{total_amount:.2f}'),
        '15 de MAYO del 2026': escape(contract_date_text),
        '15 de MAYO del 2026.': escape(contract_date_text + '.'),
        'DNI: 23260963': escape(f'DNI: {dni}'),
    }
    mapping.update(_contract_placeholder_mapping(sale))
    return mapping


def _contract_mapping_legacy(sale):
    client = sale.client
    prop = _contract_linked_property(sale)
    user = client.user if client else None
    full_name = (user.get_full_name() if user else '').strip().upper() or 'CLIENTE'
    dni = (user.dni if user else '').strip() or '00000000'
    mobile = (client.mobile if client and client.mobile else '').strip() or '000000000'
    predio_dep = (
        (prop.department if prop and prop.department else '')
        or (client.predio_department if client and client.predio_department else client.department if client else '')
        or '________________'
    )
    predio_prov = (
        (prop.province if prop and prop.province else '')
        or (client.predio_province if client and client.predio_province else client.province if client else '')
        or '________________'
    )
    predio_dist = (
        (prop.district if prop and prop.district else '')
        or (client.predio_district if client and client.predio_district else client.district if client else '')
        or '________________'
    )
    centro_poblado = (
        (prop.community if prop and prop.community else '')
        or (prop.population_center if prop and prop.population_center else '')
        or (prop.address if prop and prop.address else '')
        or (client.predio_address if client and client.predio_address else client.address if client else '')
        or '________________'
    )
    contract_date_text = _contract_signature_date_plain(sale)

    total_amount = Decimal(str(sale.total or 0)).quantize(Decimal('0.01'))
    down_payment = Decimal('0.00')
    if sale.payment_condition == 'credito':
        down_payment = Decimal(str(sale.credit_down_payment or 0)).quantize(Decimal('0.01'))
    remaining_amount = (total_amount - down_payment).quantize(Decimal('0.01'))
    if remaining_amount < 0:
        remaining_amount = Decimal('0.00')

    mapping = {
        'ELSA AROTOMA HUAMANI': escape(full_name),
        '70360638': escape(dni),
        '979769335': escape(mobile),
        'HUANCAPITE': escape(centro_poblado.upper()),
        'ANDABAMBA': escape(predio_dist.upper()),
        'ACOBAMBA': escape(predio_prov.upper()),
        'HUANCAVELICA': escape(predio_dep.upper()),
        '14 de ABRIL del 2026.': escape(contract_date_text),
        'DE S/ 700.00 SOLES': escape(f'DE S/ {total_amount:.2f} SOLES'),
        'DE S/700.00 SOLES': escape(f'DE S/ {total_amount:.2f} SOLES'),
        'S/75.00 soles': escape(f'S/ {down_payment:.2f} soles'),
        'S/ 75.00 soles': escape(f'S/ {down_payment:.2f} soles'),
        'S/625.00 soles.': escape(f'S/ {remaining_amount:.2f} soles.'),
        'S/ 625.00 soles.': escape(f'S/ {remaining_amount:.2f} soles.'),
    }
    mapping.update(_contract_placeholder_mapping(sale))
    return mapping


def _contract_mapping_from_sale(sale, template_path=None):
    if _is_celia_contract_template(template_path):
        return _contract_mapping_celia(sale)
    return _contract_mapping_legacy(sale)


def _apply_celia_location_replacements(xml_str, location_ctx):
    if not location_ctx:
        return xml_str
    dist = escape(location_ctx['district'])
    prov = escape(location_ctx['province'])
    dept = escape(location_ctx['department'])
    xml_str = xml_str.replace(' HUANCAVELICA, ', ' {} , '.format(dist))
    xml_str = xml_str.replace(' HUANCAVELICA ', ' {} '.format(prov))
    xml_str = xml_str.replace('>HUANCAVELICA</w:t>', '>' + dept + '</w:t>')
    return xml_str


def _apply_mapping_to_docx(template_bytes, mapping, sale=None, signature_date_plain=None, location_ctx=None, community=None):
    zin_mem = BytesIO(template_bytes)
    zout_mem = BytesIO()
    with ZipFile(zin_mem, 'r') as zin, ZipFile(zout_mem, 'w', ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                txt = content.decode('utf-8', 'ignore')
                if item.filename == 'word/document.xml' and signature_date_plain:
                    txt = _replace_signature_date_in_document_xml(txt, signature_date_plain)
                if item.filename == 'word/document.xml' and community:
                    txt = _replace_contract_community_in_document_xml(txt, community)
                if item.filename == 'word/document.xml' and sale:
                    txt = _replace_contract_payment_amounts_in_document_xml(txt, sale)
                for old, new in mapping.items():
                    if old:
                        txt = txt.replace(old, new)
                txt = _replace_placeholders_in_xml_paragraphs(txt, mapping)
                if location_ctx:
                    txt = _apply_celia_location_replacements(txt, location_ctx)
                content = txt.encode('utf-8')
            zout.writestr(item, content)
    return zout_mem.getvalue()


def _docx_bytes_to_preview_text(docx_bytes):
    with ZipFile(BytesIO(docx_bytes), 'r') as zf:
        xml_bytes = zf.read('word/document.xml')
    root = ET.fromstring(xml_bytes)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    lines = []
    for p in root.findall('.//w:p', ns):
        parts = []
        for t in p.findall('.//w:t', ns):
            parts.append(t.text or '')
        line = ''.join(parts).strip()
        if line:
            lines.append(line)
    return '\n'.join(lines)


def _plain_placeholder_mapping(sale):
    return {key: unescape(value) for key, value in _contract_mapping_from_sale(sale).items()}


def _iter_docx_paragraphs(container):
    for paragraph in getattr(container, 'paragraphs', []):
        yield paragraph
    for table in getattr(container, 'tables', []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_docx_paragraphs(cell)


def _replace_docx_paragraph_text(paragraph, mapping):
    text = paragraph.text or ''
    if not text:
        return False
    new_text = text
    for old, new in mapping.items():
        if old in new_text:
            new_text = new_text.replace(old, new)
    if new_text == text:
        return False
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(new_text)
    return True


def _replace_schedule_amount_lines(paragraph, sale):
    text = paragraph.text or ''
    if not text:
        return False
    total_amount, _down_payment, _remaining = _money_amounts_for_sale(sale)
    new_text = re.sub(
        r'(TOTAL\s+A\s+PAGAR\s*:\s*)S/\.\s*[_\s.]*',
        r'\1{}'.format(_format_money(total_amount)),
        text,
        flags=re.IGNORECASE,
    )
    if new_text == text:
        return False
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(new_text)
    return True


def _fill_schedule_placeholders(doc, sale):
    mapping = _plain_placeholder_mapping(sale)
    for container in [doc]:
        for paragraph in _iter_docx_paragraphs(container):
            _replace_docx_paragraph_text(paragraph, mapping)
            _replace_schedule_amount_lines(paragraph, sale)
    for section in doc.sections:
        for part_name in (
            'header', 'footer', 'first_page_header', 'first_page_footer',
            'even_page_header', 'even_page_footer',
        ):
            part = getattr(section, part_name, None)
            if part is None:
                continue
            for paragraph in _iter_docx_paragraphs(part):
                _replace_docx_paragraph_text(paragraph, mapping)
                _replace_schedule_amount_lines(paragraph, sale)


def _schedule_rows_for_sale(sale):
    ctas = CtasCollect.objects.filter(sale=sale).order_by('id').first()
    if ctas:
        return ctas.get_quota_plan()

    total, inicial, financed = _money_amounts_for_sale(sale)
    n = max(1, int(sale.credit_quota_count or 1))
    base_date = sale.date_joined
    end_date = sale.end_credit
    diff_days = (end_date - base_date).days if end_date and base_date else 25 * n
    if diff_days <= 0:
        diff_days = 25 * n
    amounts = []
    cents = int((financed * 100).to_integral_value())
    q, r = divmod(cents, n)
    for i in range(n):
        amounts.append(Decimal(q + (1 if i < r else 0)) / Decimal(100))

    rows = []
    if inicial > 0:
        rows.append({
            'num': 0,
            'label': 'Inicial',
            'amount': format(inicial, '.2f'),
            'due_date': base_date.strftime('%Y-%m-%d'),
        })
    from datetime import timedelta
    for i in range(1, n + 1):
        due = base_date + timedelta(days=round((diff_days * i) / n))
        rows.append({
            'num': i,
            'label': 'Cuota {}'.format(i),
            'amount': format(amounts[i - 1], '.2f'),
            'due_date': due.strftime('%Y-%m-%d'),
        })
    return rows


def _table_text(table):
    parts = []
    for row in table.rows:
        for cell in row.cells:
            parts.append(cell.text or '')
    return ' '.join(parts).lower()


def _find_schedule_table(doc):
    candidates = []
    for table in doc.tables:
        text = _table_text(table)
        score = 0
        for token in ('cuota', 'fecha', 'monto', 'pendiente'):
            if token in text:
                score += 1
        if score >= 2:
            candidates.append((score, table))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _clear_cell(cell):
    for paragraph in cell.paragraphs:
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        if paragraph.text:
            paragraph.text = ''


def _set_cell_text(cell, value):
    _clear_cell(cell)
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(value))
    else:
        cell.text = str(value)


def _schedule_header_values(sale):
    mapping = _plain_placeholder_mapping(sale)
    predio_name = (
        mapping.get('{{PREDIO_TIPO}}', '')
        or mapping.get('{{SERVICIO}}', '')
        or mapping.get('{{PRODUCTO}}', '')
    )
    return {
        'nombre': mapping.get('{{CLIENTE_NOMBRE}}', ''),
        'dni': mapping.get('{{CLIENTE_DNI}}', ''),
        'predio': predio_name,
        'comunidad': mapping.get('{{PREDIO_COMUNIDAD}}', ''),
        'direccion': mapping.get('{{PREDIO_DIRECCION}}', ''),
    }


def _fill_schedule_header_tables(doc, sale):
    values = _schedule_header_values(sale)
    changed = False
    label_map = (
        ('nombre y apellido', values['nombre']),
        ('dni', values['dni']),
        ('predio', values['predio']),
        ('nombre de comunidad', values['comunidad']),
        ('direccion del predio', values['direccion']),
        ('dirección del predio', values['direccion']),
    )
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = (row.cells[0].text or '').strip().lower().replace(':', '')
            normalized = (
                label.replace('á', 'a')
                .replace('é', 'e')
                .replace('í', 'i')
                .replace('ó', 'o')
                .replace('ú', 'u')
            )
            for key, value in label_map:
                key_norm = (
                    key.replace('á', 'a')
                    .replace('é', 'e')
                    .replace('í', 'i')
                    .replace('ó', 'o')
                    .replace('ú', 'u')
                )
                if normalized == key_norm and value:
                    _set_cell_text(row.cells[1], value)
                    changed = True
                    break
    return changed


def _fill_existing_schedule_table(doc, sale):
    table = _find_schedule_table(doc)
    if table is None:
        return False
    rows = _schedule_rows_for_sale(sale)
    if not rows:
        return True

    header_idx = 0
    for idx, row in enumerate(table.rows):
        row_text = ' '.join(cell.text for cell in row.cells).lower()
        if 'cuota' in row_text and ('fecha' in row_text or 'monto' in row_text):
            header_idx = idx
            break
    first_data_idx = header_idx + 1
    while len(table.rows) < first_data_idx + len(rows):
        table.add_row()

    for idx, item in enumerate(rows):
        row = table.rows[first_data_idx + idx]
        cells = row.cells
        if len(cells) < 3:
            continue
        _set_cell_text(cells[0], item.get('label') or '')
        _set_cell_text(cells[1], item.get('due_date') or '')
        _set_cell_text(cells[2], 'S/ {}'.format(item.get('amount') or '0.00'))
        if len(cells) >= 4:
            _set_cell_text(cells[3], 'Pendiente' if item.get('num') else 'Pagado')
        if len(cells) >= 5:
            _set_cell_text(cells[4], '')

    for idx in range(first_data_idx + len(rows), len(table.rows)):
        for cell in table.rows[idx].cells:
            _set_cell_text(cell, '')
    return True


def _append_schedule_table(doc, sale):
    rows = _schedule_rows_for_sale(sale)
    target_paragraph = None
    marker = '{{CRONOGRAMA_TABLA}}'
    for paragraph in _iter_docx_paragraphs(doc):
        if marker in (paragraph.text or ''):
            target_paragraph = paragraph
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            break

    if target_paragraph is None:
        doc.add_paragraph('Cronograma de pagos')

    table = doc.add_table(rows=1, cols=3)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    headers = table.rows[0].cells
    headers[0].text = 'Cuota'
    headers[1].text = 'Fecha programada'
    headers[2].text = 'Monto'
    for item in rows:
        cells = table.add_row().cells
        cells[0].text = item.get('label') or ''
        cells[1].text = item.get('due_date') or ''
        cells[2].text = 'S/ {}'.format(item.get('amount') or '0.00')

    if target_paragraph is not None:
        target_paragraph._p.addnext(table._tbl)


def _build_payment_schedule_docx(template_bytes, sale):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError('Falta python-docx para generar el cronograma de pagos.')

    try:
        doc = Document(BytesIO(template_bytes))
        _fill_schedule_placeholders(doc, sale)
        _fill_schedule_header_tables(doc, sale)
    except Exception:
        doc = Document()
        doc.add_heading('Cronograma de pagos', level=1)
        mapping = _plain_placeholder_mapping(sale)
        client_name = mapping.get('{{CLIENTE_NOMBRE}}', 'CLIENTE')
        doc.add_paragraph('Cliente: {}'.format(client_name))
        doc.add_paragraph('Contrato: {}'.format(mapping.get('{{CODIGO_CONTRATO}}', '')))
        doc.add_paragraph('Monto total: {}'.format(mapping.get('{{MONTO_TOTAL}}', '')))
    if not _fill_existing_schedule_table(doc, sale):
        _append_schedule_table(doc, sale)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _archive_contract_docx(sale, docx_bytes):
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    archive_dir = os.path.join(settings.MEDIA_ROOT, 'contracts_archive')
    os.makedirs(archive_dir, exist_ok=True)
    cc = (sale.contract_code or f'CT-{sale.id:06d}').replace('/', '-').replace('\\', '-')
    file_name = f'{cc}_{ts}.docx'
    file_path = os.path.join(archive_dir, file_name)
    with open(file_path, 'wb') as f:
        f.write(docx_bytes)
    return file_path


DOCX_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
)

_SALE_TICKET_TEMPLATES = {
    'ticket': 'crm/sale/print/ticket.html',
    'ticket-rppos': 'crm/sale/print/ticket_rppos.html',
    'ticket-58': 'crm/sale/print/ticket_rppos.html',
    'ticket-termica': 'crm/sale/print/ticket_termica.html',
    'ticket-80': 'crm/sale/print/ticket_termica.html',
}


class SalePrintVoucherView(LoginRequiredMixin, View):
    success_url = reverse_lazy('sale_admin_list')

    def get_success_url(self):
        if self.request.user.is_client():
            return reverse_lazy('sale_client_list')
        return self.success_url

    def get_height_ticket(self):
        sale = Sale.objects.get(pk=self.kwargs['pk'])
        height = 115
        increment = sale.saledetail_set.all().count() * 5.45
        height += increment
        return round(height)

    def get_height_ticket_80(self):
        sale = Sale.objects.get(pk=self.kwargs['pk'])
        height = 135
        increment = sale.saledetail_set.all().count() * 4.2
        height += increment
        return round(height)

    def get(self, request, *args, **kwargs):
        try:
            sale = Sale.objects.get(pk=self.kwargs['pk'])
            voucher = (self.kwargs.get('voucher') or 'ticket').lower()
            print_ctx = comprobante_print_context()
            if request.GET.get('format', 'html').lower() != 'pdf':
                print_ctx = dict(print_ctx)
                print_ctx['comprobante_logo'] = request.build_absolute_uri(comprobante_logo_url())
            context = {
                'sale': sale,
                'company': Company.objects.first(),
                **print_ctx,
            }
            if sale.type_voucher == 'ticket':
                ticket_template = _SALE_TICKET_TEMPLATES.get(
                    voucher, _SALE_TICKET_TEMPLATES['ticket'],
                )
                template = get_template(ticket_template)
                if voucher in ('ticket-termica', 'ticket-80'):
                    context['height'] = self.get_height_ticket_80()
                else:
                    context['height'] = self.get_height_ticket()
            else:
                template = get_template('crm/sale/print/invoice.html')
            html_content = template.render(context)
            if request.GET.get('format', 'pdf').lower() == 'html':
                return HttpResponse(html_content, content_type='text/html; charset=utf-8')
            url_css = os.path.join(
                settings.BASE_DIR, 'static/lib/bootstrap-4.6.0/css/bootstrap.min.css',
            )
            pdf_file = HTML(
                string=html_content.encode(encoding='UTF-8'),
                base_url=request.build_absolute_uri(),
            ).write_pdf(
                stylesheets=[CSS(url_css)], presentational_hints=True,
            )
            return HttpResponse(pdf_file, content_type='application/pdf')
        except Exception as exc:
            return HttpResponse(
                'Error al generar el comprobante: {}'.format(exc),
                status=500,
                content_type='text/plain; charset=utf-8',
            )


class SalePrintContractView(LoginRequiredMixin, View):
    success_url = reverse_lazy('sale_admin_list')
    raise_exception = True

    def get_success_url(self):
        if self.request.user.is_client():
            return reverse_lazy('sale_client_list')
        return self.success_url

    def get(self, request, *args, **kwargs):
        template_path = _contract_template_path()
        if not os.path.isfile(template_path):
            return _contract_download_error_response(
                request,
                'No existe la plantilla del contrato (contrato_Celia.docx o contrato_base.docx).',
                status=500,
            )
        try:
            sale = Sale.objects.select_related('client__user').get(pk=self.kwargs['pk'])
            mapping = _contract_mapping_from_sale(sale, template_path)
            with open(template_path, 'rb') as f:
                template_bytes = f.read()
            sig_date = (
                _contract_signature_date_celia(sale)
                if _is_celia_contract_template(template_path)
                else _contract_signature_date_plain(sale)
            )
            docx_out = _build_decorated_contract_docx(
                template_bytes,
                sale,
                mapping,
                sig_date,
                template_path=template_path,
            )
            out_name = sale.contract_docx_basename()
            response = HttpResponse(docx_out, content_type=DOCX_CONTENT_TYPE)
            response['Content-Disposition'] = f'attachment; filename="{out_name}"'
            response['X-Contract-Filename'] = out_name
            lock_client_properties_for_sale_contract(sale)
            return response
        except Sale.DoesNotExist:
            return _contract_download_error_response(request, 'Venta no encontrada.', status=404)
        except Exception:
            return _contract_download_error_response(
                request,
                'Error al generar el contrato. Revise datos del cliente y la plantilla Word.',
                status=500,
            )


class SalePrintPaymentScheduleView(LoginRequiredMixin, View):
    success_url = reverse_lazy('sale_admin_list')
    raise_exception = True

    def get_success_url(self):
        if self.request.user.is_client():
            return reverse_lazy('sale_client_list')
        return self.success_url

    def get(self, request, *args, **kwargs):
        template_path = _schedule_template_path()
        if not os.path.isfile(template_path):
            return _contract_download_error_response(
                request,
                'No existe la plantilla del cronograma de pagos (Cronograma.docx).',
                status=500,
            )
        try:
            sale = Sale.objects.select_related('client__user', 'employee').get(pk=self.kwargs['pk'])
            if sale.payment_condition != 'credito':
                return HttpResponse('', status=204)
            with open(template_path, 'rb') as f:
                template_bytes = f.read()
            docx_out = _build_payment_schedule_docx(template_bytes, sale)
            docx_out = decorate_contract_celia_docx(
                docx_out,
                _client_code_for_sale(sale),
                _sale_code_for_sale(sale),
                _contract_code_for_documents(sale),
            )
            sale_code = _sale_code_for_sale(sale).replace('/', '-').replace('\\', '-')
            out_name = 'CRONOGRAMA_PAGOS_{}.docx'.format(sale_code or sale.pk)
            response = HttpResponse(docx_out, content_type=DOCX_CONTENT_TYPE)
            response['Content-Disposition'] = f'attachment; filename="{out_name}"'
            response['X-Schedule-Filename'] = out_name
            return response
        except Sale.DoesNotExist:
            return _contract_download_error_response(request, 'Venta no encontrada.', status=404)
        except Exception:
            return _contract_download_error_response(
                request,
                'Error al generar el cronograma de pagos. Revise datos del crédito y la plantilla Word.',
                status=500,
            )


class SalePrintContractPreviewView(LoginRequiredMixin, View):
    success_url = reverse_lazy('sale_admin_list')

    def get_success_url(self):
        if self.request.user.is_client():
            return reverse_lazy('sale_client_list')
        return self.success_url

    def get(self, request, *args, **kwargs):
        try:
            sale = Sale.objects.select_related('client__user').get(pk=self.kwargs['pk'])
            template_path = _contract_template_path()
            mapping = _contract_mapping_from_sale(sale, template_path)
            with open(template_path, 'rb') as f:
                template_bytes = f.read()
            sig_date = (
                _contract_signature_date_celia(sale)
                if _is_celia_contract_template(template_path)
                else _contract_signature_date_plain(sale)
            )
            docx_out = _build_decorated_contract_docx(
                template_bytes,
                sale,
                mapping,
                sig_date,
                template_path=template_path,
            )
            preview_text = _docx_bytes_to_preview_text(docx_out)
            return render(
                request,
                'crm/sale/print/contract_preview.html',
                {
                    'sale': sale,
                    'preview_text': preview_text,
                    'embed': request.GET.get('embed') == '1',
                },
            )
        except Exception:
            return HttpResponseRedirect(self.get_success_url())


class SalePrintContractQuickView(LoginRequiredMixin, View):
    success_url = reverse_lazy('sale_admin_list')

    def get_success_url(self):
        if self.request.user.is_client():
            return reverse_lazy('sale_client_list')
        return self.success_url

    def get(self, request, *args, **kwargs):
        try:
            sale = Sale.objects.select_related('client__user').get(pk=self.kwargs['pk'])
            template_path = _contract_template_path()
            mapping = _contract_mapping_from_sale(sale, template_path)
            with open(template_path, 'rb') as f:
                template_bytes = f.read()
            sig_date = (
                _contract_signature_date_celia(sale)
                if _is_celia_contract_template(template_path)
                else _contract_signature_date_plain(sale)
            )
            docx_out = _build_decorated_contract_docx(
                template_bytes,
                sale,
                mapping,
                sig_date,
                template_path=template_path,
            )
            _archive_contract_docx(sale, docx_out)
            lock_client_properties_for_sale_contract(sale)
            preview_text = _docx_bytes_to_preview_text(docx_out)
            return render(
                request,
                'crm/sale/print/contract_quick_print.html',
                {'sale': sale, 'preview_text': preview_text},
            )
        except Exception:
            return HttpResponseRedirect(self.get_success_url())
