"""
Decoración del contrato Word (.docx): código de barras, URL empresa, código QR y leyenda breve bajo el QR.

Hay dos modos (ver PLACEHOLDERS_CONTRATO.txt junto a la plantilla):

1) Plantilla con marcadores Polariss — respeta posición y formato que definas en Word:
   [[POLARISS_BARCODE]]  → barras + código de contrato en ese párrafo
   [[POLARISS_WEB]]      → URL empresa en ese párrafo (alineación la defines tú)
   [[POLARISS_QR]]       → QR + código CT + línea breve sobre la empresa (descripción Empresa o texto por defecto)

   Las variantes antiguas [[FACTORA_*]] siguen reconocidas por compatibilidad.

   Si aparece al menos un marcador en el .docx fusionado, SOLO se sustituyen esos párrafos.
   Logo, encabezados, pies y el resto del formato de la plantilla no se tocan.

2) Sin marcadores POLARISS_* / FACTORA_* — no se inserta código de barras, URL ni QR; el archivo fusionado se devuelve igual (conserva diseño completo).
"""
from __future__ import annotations

import logging
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)

try:
    import barcode
    from barcode.writer import ImageWriter
except ImportError:
    barcode = None  # type: ignore

try:
    import qrcode
except ImportError:
    qrcode = None  # type: ignore

try:
    from PIL import Image as PILImage  # noqa: F401 — comprueba Pillow para barcode/qrcode
except ImportError:
    PILImage = None  # type: ignore

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:
    Document = None  # type: ignore


# Marcadores en la plantilla Word (un marcador por párrafo; ver PLACEHOLDERS_CONTRATO.txt).
MARK_BARCODE = '[[POLARISS_BARCODE]]'
MARK_WEB = '[[POLARISS_WEB]]'
MARK_QR = '[[POLARISS_QR]]'
# Plantillas antiguas con prefijo FACTORA (se siguen sustituyendo).
MARK_BARCODE_LEGACY = '[[FACTORA_BARCODE]]'
MARK_WEB_LEGACY = '[[FACTORA_WEB]]'
MARK_QR_LEGACY = '[[FACTORA_QR]]'
ALL_MARKERS = (
    MARK_BARCODE,
    MARK_WEB,
    MARK_QR,
    MARK_BARCODE_LEGACY,
    MARK_WEB_LEGACY,
    MARK_QR_LEGACY,
)

# Texto bajo el QR si la empresa no tiene campo «descripción» en BD (saneamiento predial).
DEFAULT_QR_COMPANY_LEGEND = 'Empresa dedicada al saneamiento predial.'


def decoration_deps_available():
    """True si hay motor suficiente para intentar decorar."""
    return Document is not None


def _clear_paragraph_runs(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def iter_body_paragraphs(parent):
    if parent is None:
        return
    for p in parent.paragraphs:
        yield p
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_body_paragraphs(cell)


def iter_all_paragraphs(doc):
    """Párrafos del cuerpo y de cabeceras/pies (sin duplicar el mismo nodo XML si hay secciones enlazadas)."""
    seen = set()

    def gen(part):
        if part is None:
            return
        for p in iter_body_paragraphs(part):
            key = id(p._element)
            if key in seen:
                continue
            seen.add(key)
            yield p

    yield from gen(doc)
    for section in doc.sections:
        for name in (
            'header',
            'footer',
            'first_page_header',
            'first_page_footer',
            'even_page_header',
            'even_page_footer',
        ):
            if not hasattr(section, name):
                continue
            try:
                yield from gen(getattr(section, name))
            except Exception:
                continue


def _document_has_placeholders(doc) -> bool:
    for p in iter_all_paragraphs(doc):
        t = p.text or ''
        for m in ALL_MARKERS:
            if m in t:
                return True
    return False


def _fill_barcode_placeholder(paragraph, barcode_png: Optional[BytesIO], contract_code: str):
    _clear_paragraph_runs(paragraph)
    if barcode_png is not None:
        barcode_png.seek(0)
        paragraph.add_run().add_picture(barcode_png, width=Inches(2.25), height=Inches(0.5))
        paragraph.add_run(' ')
    r = paragraph.add_run(contract_code)
    r.bold = True


def _fill_web_placeholder(paragraph, url_show: str):
    _clear_paragraph_runs(paragraph)
    paragraph.add_run(url_show or '')


def _fill_qr_placeholder(paragraph, qr_png: Optional[BytesIO], contract_code: str, qr_legend: str):
    _clear_paragraph_runs(paragraph)
    # Alineación y sangría: las define el párrafo en la plantilla Word.
    if qr_png is not None:
        qr_png.seek(0)
        paragraph.add_run().add_picture(qr_png, width=Inches(1.35))
    paragraph.add_run('\n')
    paragraph.add_run(contract_code).font.size = Pt(9)
    leg = (qr_legend or '').strip()
    if leg:
        paragraph.add_run('\n')
        lr = paragraph.add_run(leg)
        lr.font.size = Pt(8)


def _apply_placeholder_decorations(
    doc,
    cc: str,
    url_show: str,
    bc_png: Optional[BytesIO],
    qr_png: Optional[BytesIO],
    qr_legend: str,
) -> None:
    counts = {'barcode': 0, 'web': 0, 'qr': 0}
    for p in iter_all_paragraphs(doc):
        txt = (p.text or '').strip()
        if txt in (MARK_BARCODE, MARK_BARCODE_LEGACY):
            _fill_barcode_placeholder(p, bc_png, cc)
            counts['barcode'] += 1
        elif txt in (MARK_WEB, MARK_WEB_LEGACY):
            _fill_web_placeholder(p, url_show)
            counts['web'] += 1
        elif txt in (MARK_QR, MARK_QR_LEGACY):
            _fill_qr_placeholder(p, qr_png, cc, qr_legend)
            counts['qr'] += 1
    logger.info('Contrato sustituyendo marcadores Polariss / FACTORA (legacy): %s', counts)


def resize_png_stream(png_stream, width, height):
    """Ajusta barcode/QR al tamaño de la imagen en la plantilla (evita tapar sello/firma)."""
    if PILImage is None or not width or not height:
        png_stream.seek(0)
        return png_stream.read()
    try:
        png_stream.seek(0)
        img = PILImage.open(png_stream).convert('RGBA')
        img = img.resize((int(width), int(height)), PILImage.Resampling.LANCZOS)
        buf = BytesIO()
        img.save(buf, format='PNG')
        return buf.getvalue()
    except Exception as exc:
        logger.warning('No se redimensionó PNG de plantilla: %s', exc)
        png_stream.seek(0)
        return png_stream.read()


def _generate_barcode_png(contract_code: str) -> Optional[BytesIO]:
    """
    Code128 con el código de contrato completo (ej. CT-000042).
    Se conserva el guión para que la pistola lectora devuelva exactamente el código.
    Parámetros optimizados para lectura con escáner físico.
    """
    if not barcode or PILImage is None or not contract_code:
        return None
    safe = contract_code.strip()
    if not safe:
        return None
    try:
        writer = ImageWriter()
        bc_cls = barcode.get_barcode_class('code128')
        bc = bc_cls(safe, writer=writer)
        buf = BytesIO()
        try:
            bc.write(
                buf,
                options={
                    'module_width': 0.33,
                    'module_height': 12,
                    'quiet_zone': 6.5,
                    'write_text': True,
                    'font_size': 10,
                    'text_distance': 3,
                    'dpi': 300,
                },
            )
        except TypeError:
            buf = BytesIO()
            bc.write(buf)
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.warning('No se generó código de barras (%s): %s', safe, exc)
        return None


def _generate_qr_png(contract_code: str) -> Optional[BytesIO]:
    """QR con el código de contrato completo (ej. CT-000042), tamaño optimizado para impresión."""
    if not qrcode or PILImage is None or not contract_code:
        return None
    safe = contract_code.strip()
    if not safe:
        return None
    try:
        qr = qrcode.QRCode(version=None, box_size=8, border=3, error_correction=qrcode.constants.ERROR_CORRECT_M)
        qr.add_data(safe)
        qr.make(fit=True)
        img = qr.make_image(fill_color='black', back_color='white')
        buf = BytesIO()
        img.save(buf, format='PNG')
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.warning('No se generó QR (%s): %s', safe, exc)
        return None


def _normalize_public_url(website: Optional[str]) -> str:
    if not website or not str(website).strip():
        return ''
    w = str(website).strip()
    if '://' not in w:
        return 'https://' + w.lstrip('/')
    return w


def _resolve_qr_legend(company_description: Optional[str]) -> str:
    s = (company_description or '').strip()
    return s if s else DEFAULT_QR_COMPANY_LEGEND


def _replace_images_in_docx_zip(docx_bytes: bytes, contract_code: str) -> bytes:
    """
    Reemplaza directamente las imágenes del body del documento que corresponden
    al código de barras y QR estáticos por los generados con el código de contrato.
    Identifica por dimensiones: barcode = ancho > 2*alto; QR = relación ~1:1 y <= 200x200.
    Solo toca imágenes referenciadas desde word/document.xml (no headers/footers/logo).
    """
    import struct
    from zipfile import ZipFile, ZIP_DEFLATED

    if not barcode and not qrcode:
        return docx_bytes

    cc = (contract_code or '').strip()
    if not cc:
        return docx_bytes

    zin_mem = BytesIO(docx_bytes)
    zout_mem = BytesIO()

    try:
        with ZipFile(zin_mem, 'r') as zin:
            # Read document.xml.rels to find body image references
            rels_path = 'word/_rels/document.xml.rels'
            if rels_path not in zin.namelist():
                return docx_bytes

            import xml.etree.ElementTree as ET
            rels_xml = zin.read(rels_path).decode('utf-8', 'ignore')
            rels_root = ET.fromstring(rels_xml)
            body_images = {}
            for rel in rels_root:
                rel_type = rel.get('Type', '')
                target = rel.get('Target', '')
                if 'image' in rel_type.lower() and target:
                    full_path = 'word/' + target if not target.startswith('word/') else target
                    body_images[full_path] = rel.get('Id', '')

            # Classify images by dimensions
            barcode_path = None
            qr_path = None
            for img_path in body_images:
                if img_path not in zin.namelist():
                    continue
                data = zin.read(img_path)
                if len(data) < 24 or data[:4] != b'\x89PNG':
                    continue
                w = struct.unpack('>I', data[16:20])[0]
                h = struct.unpack('>I', data[20:24])[0]
                if w > 2 * h and w > 200:
                    barcode_path = img_path
                elif abs(w - h) < max(w, h) * 0.3 and w <= 250 and h <= 250:
                    qr_path = img_path

            if not barcode_path and not qr_path:
                return docx_bytes

            bc_png = _generate_barcode_png(cc) if barcode_path else None
            qr_png = _generate_qr_png(cc) if qr_path else None

            with ZipFile(zout_mem, 'w', ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    content = zin.read(item.filename)
                    if item.filename == barcode_path and bc_png is not None:
                        bc_png.seek(0)
                        content = bc_png.read()
                    elif item.filename == qr_path and qr_png is not None:
                        qr_png.seek(0)
                        content = qr_png.read()
                    zout.writestr(item, content)

        return zout_mem.getvalue()
    except Exception as exc:
        logger.warning('Fallo reemplazo de imágenes barcode/QR: %s', exc, exc_info=True)
        return docx_bytes


def decorate_contract_docx(
    docx_bytes: bytes,
    contract_code: str,
    company_website: Optional[str],
    company_description: Optional[str] = None,
) -> bytes:
    """
    Modo 1: Si la plantilla tiene marcadores [[POLARISS_*]] / [[FACTORA_*]], los sustituye
    por código de barras, URL y QR generados con el código de contrato.

    Modo 2: Si NO hay marcadores, busca imágenes estáticas de barcode/QR en el body
    y las reemplaza in-situ con los códigos generados para esta venta.
    """
    if Document is None:
        logger.error(
            'Contrato sin código QR/barcode en Word: instale dependencias '
            '(python-docx, qrcode, python-barcode, Pillow). Ej.: pip install python-docx qrcode python-barcode Pillow'
        )
        return docx_bytes

    if PILImage is None:
        logger.error(
            'Contrato sin imágenes barcode/QR: falta Pillow. Ej.: pip install Pillow'
        )

    cc = (contract_code or '').strip()
    if not cc:
        return docx_bytes

    url_show = _normalize_public_url(company_website)
    qr_legend = _resolve_qr_legend(company_description)

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as exc:
        logger.exception('No se pudo abrir el .docx del contrato con python-docx: %s', exc)
        return docx_bytes

    if not _document_has_placeholders(doc):
        logger.info(
            'Contrato Word sin marcadores POLARISS_*/FACTORA_*: '
            'intentando reemplazar imágenes estáticas de barcode/QR por código dinámico.'
        )
        return _replace_images_in_docx_zip(docx_bytes, cc)

    bc_png = _generate_barcode_png(cc)
    qr_png = _generate_qr_png(cc)

    try:
        _apply_placeholder_decorations(doc, cc, url_show, bc_png, qr_png, qr_legend)
    except Exception as exc:
        logger.warning('Fallo decoración contrato Word: %s', exc, exc_info=True)

    try:
        out = BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as exc:
        logger.exception('No se pudo guardar el .docx decorado: %s', exc)
        return docx_bytes
