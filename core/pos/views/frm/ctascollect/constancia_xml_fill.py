"""
Relleno de plantilla constancia preservando runs, imágenes y formato del Word.
"""
from __future__ import annotations

import re

from core.pos.views.crm.sale.print.docx_codes import iter_body_paragraphs
from core.pos.views.crm.sale.print.views import W_NS, W_P, _paragraph_join_text

try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    WD_ALIGN_PARAGRAPH = None

_ELLIPSIS_RE = re.compile(r'^[\s\u2026.]+$')
# Bloques de puntos suspensivos del Word (……) dentro del mismo run
_INLINE_PLACEHOLDER_RE = re.compile(r'(?:[\u2026.]){2,}')


def _is_ellipsis_text(text):
    if not text or not str(text).strip():
        return False
    s = str(text).strip()
    if len(s) < 2:
        return False
    if _ELLIPSIS_RE.match(s):
        return True
    # Puntos o guiones de relleno del Word (. . . o ……)
    if all(c in '.\u2026 ' for c in s) and len(s) >= 2:
        return True
    return False


def _paragraph_has_placeholders(p_elem):
    for t_elem in p_elem.findall('.//{%s}t' % W_NS):
        if t_elem.text and _is_ellipsis_text(t_elem.text):
            return True
    return False


def _replace_ellipsis_runs_in_order(p_elem, values):
    """Sustituye bloques … o .. en orden (en el mismo w:t o en runs distintos)."""
    if not values:
        return
    idx = 0
    for t_elem in p_elem.findall('.//{%s}t' % W_NS):
        if idx >= len(values) or not t_elem.text:
            continue
        text = t_elem.text
        if _is_ellipsis_text(text):
            t_elem.text = values[idx]
            idx += 1
            continue
        while idx < len(values):
            m = _INLINE_PLACEHOLDER_RE.search(text)
            if not m:
                break
            text = text[: m.start()] + values[idx] + text[m.end() :]
            idx += 1
        t_elem.text = text


def _is_company_signature_paragraph(text):
    t = (text or '').strip().upper()
    if not t:
        return False
    return (
        'TERRAFIRMA' in t
        or t.startswith('G.G.')
        or 'ASESOR' in t
        or 'CONSULTOR' in t
    )


def _set_cell_centered_lines(cell, lines):
    cell.text = ''
    for idx, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if WD_ALIGN_PARAGRAPH is not None:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)


def ensure_constancia_signature_block(doc, ctx):
    """Reemplaza la firma única de la plantilla por cliente (izq.) y empresa (der.)."""
    anchor = None
    to_remove = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith('________________'):
            if anchor is None:
                anchor = paragraph
            to_remove.append(paragraph)
            continue
        if anchor is not None and _is_company_signature_paragraph(text):
            to_remove.append(paragraph)
            break

    if anchor is None:
        return

    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass

    client_name = (ctx.get('name') or 'CLIENTE').strip()
    rep_name = (ctx.get('rep_name') or '').strip()
    company_name = (ctx.get('company_name') or 'TERRAFIRMA ASESORÍA Y CONSULTORÍA').strip()

    _set_cell_centered_lines(
        table.rows[0].cells[0],
        ['_' * 28, 'Firma del cliente', client_name],
    )
    _set_cell_centered_lines(
        table.rows[0].cells[1],
        ['_' * 28, rep_name, company_name],
    )

    anchor._element.addprevious(table._tbl)
    for paragraph in to_remove:
        paragraph._element.getparent().remove(paragraph._element)


def _replace_signature_date_runs(p_elem, ctx):
    """Fecha al final: día, mes y año en los tres bloques … del párrafo."""
    parts = ctx.get('signature_date_parts')
    if not parts:
        parts = [ctx.get('signature_date', '')]
    _replace_ellipsis_runs_in_order(p_elem, parts)


def fill_constancia_document_xml(xml_str, ctx):
    import xml.etree.ElementTree as ET

    try:
        root = ET.fromstring(xml_str.encode('utf-8'))
    except ET.ParseError:
        return xml_str

    predio_tail = (
        'titulación de predio {}:'.format(ctx['predio_desc'])
        if ctx.get('predio_desc')
        else 'titulación de predio:'
    )

    for p_elem in root.iter(W_P):
        full = _paragraph_join_text(p_elem)
        if not full.strip():
            continue
        if 'Por medio de la presente' in full:
            _replace_ellipsis_runs_in_order(
                p_elem,
                [
                    ctx['name'],
                    ctx['dni'],
                    ctx['address'],
                    ctx['province'],
                    ctx['department'],
                ],
            )
            # Texto fijo del párrafo: actualizar solo el tramo del predio si hay descripción
            if ctx.get('predio_desc'):
                for t_elem in p_elem.findall('.//{%s}t' % W_NS):
                    if t_elem.text and 'titulación de predio' in t_elem.text:
                        t_elem.text = predio_tail
                        break
        elif 'monto total de' in full and (
            'soles' in full.lower() or _paragraph_has_placeholders(p_elem)
        ):
            _replace_ellipsis_runs_in_order(
                p_elem,
                [ctx['total_fmt'], ctx['total_words']],
            )
        elif 'CUOTA' in full.upper() and (
            'soles' in full.lower() or _paragraph_has_placeholders(p_elem)
        ):
            _replace_ellipsis_runs_in_order(
                p_elem,
                [ctx['quota_label'], ctx['quota_amount_fmt'], ctx['quota_words']],
            )
        elif 'El presente pago se realiz' in full:
            _replace_ellipsis_runs_in_order(
                p_elem,
                [ctx['payment_method'], ctx.get('collector_name') or ctx.get('beneficiary', '')],
            )
        elif 'conformidad' in full and 'HUANCAVELICA' in full.upper():
            _replace_signature_date_runs(p_elem, ctx)

    signature_targets = []
    for p_elem in root.iter(W_P):
        full = _paragraph_join_text(p_elem)
        t = full.strip()
        if not t:
            continue
        if t.startswith('________________'):
            signature_targets.append(p_elem)
        elif t and _is_ellipsis_text(t) and len(t) > 5:
            signature_targets.append(p_elem)
    if len(signature_targets) >= 2:
        _replace_ellipsis_runs_in_order(signature_targets[-2], [ctx['rep_name']])
        _replace_ellipsis_runs_in_order(signature_targets[-1], [ctx['rep_title']])
    elif len(signature_targets) == 1:
        _replace_ellipsis_runs_in_order(signature_targets[0], [ctx['rep_name']])

    ET.register_namespace('w', W_NS)
    return ET.tostring(root, encoding='unicode', method='xml')


def _paragraph_xml(paragraph):
    if hasattr(paragraph, '_element'):
        return paragraph._element
    return paragraph


def fill_constancia_doc_paragraphs(doc, ctx):
    """Rellena solo campos … en el cuerpo (no cabecera/pie); no reescribe el XML del .docx."""
    predio_tail = (
        'titulación de predio {}:'.format(ctx['predio_desc'])
        if ctx.get('predio_desc')
        else 'titulación de predio:'
    )

    for paragraph in iter_body_paragraphs(doc):
        p_elem = _paragraph_xml(paragraph)
        full = _paragraph_join_text(p_elem)
        if not full.strip():
            continue
        if 'Por medio de la presente' in full:
            _replace_ellipsis_runs_in_order(
                p_elem,
                [
                    ctx['name'],
                    ctx['dni'],
                    ctx['address'],
                    ctx['province'],
                    ctx['department'],
                ],
            )
            if ctx.get('predio_desc'):
                for t_elem in p_elem.findall('.//{%s}t' % W_NS):
                    if t_elem.text and 'titulación de predio' in t_elem.text:
                        t_elem.text = predio_tail
                        break
        elif 'monto total de' in full and (
            'soles' in full.lower() or _paragraph_has_placeholders(p_elem)
        ):
            _replace_ellipsis_runs_in_order(
                p_elem,
                [ctx['total_fmt'], ctx['total_words']],
            )
        elif 'CUOTA' in full.upper() and (
            'soles' in full.lower() or _paragraph_has_placeholders(p_elem)
        ):
            _replace_ellipsis_runs_in_order(
                p_elem,
                [ctx['quota_label'], ctx['quota_amount_fmt'], ctx['quota_words']],
            )
        elif 'El presente pago se realiz' in full:
            _replace_ellipsis_runs_in_order(
                p_elem,
                [ctx['payment_method'], ctx.get('collector_name') or ctx.get('beneficiary', '')],
            )
        elif 'conformidad' in full and 'HUANCAVELICA' in full.upper():
            _replace_signature_date_runs(p_elem, ctx)

    ensure_constancia_signature_block(doc, ctx)
